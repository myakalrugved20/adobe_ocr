"""
Generic Word document builder.
Takes structured page/block/image data and produces a .docx file.
"""

import os
import re
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from backend.config import PT_TO_EMU
from backend.services.pdf_extractor import _is_inside_table

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'v': 'urn:schemas-microsoft-com:vml',
    'w10': 'urn:schemas-microsoft-com:office:word',
}


def _vqn(tag):
    """Resolve namespace-prefixed tag using our NSMAP (supports v: and w10: that python-docx qn() doesn't)."""
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'

# Regex to strip characters invalid in XML 1.0
_INVALID_XML_RE = re.compile(
    r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]'
)


def _safe(text):
    """Ensure text is a string with only XML-valid characters."""
    if not isinstance(text, str):
        text = str(text) if text is not None else ''
    return _INVALID_XML_RE.sub('', text)


def _char_width_multiplier(text):
    """Return a per-character width multiplier based on the script of the text.

    Latin/ASCII chars are narrow (~0.6× font size), while Devanagari, CJK,
    Arabic, and other complex scripts are wider (~0.85-1.0× font size).
    """
    if not text:
        return 0.6
    wide = 0
    total = 0
    for ch in text:
        if ch.isspace():
            continue
        total += 1
        cp = ord(ch)
        # Devanagari (Hindi, Marathi, Sanskrit, etc.)
        if 0x0900 <= cp <= 0x097F or 0xA8E0 <= cp <= 0xA8FF:
            wide += 1
        # CJK Unified Ideographs
        elif 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF:
            wide += 1
        # Arabic
        elif 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F:
            wide += 1
        # Bengali, Gujarati, Tamil, Telugu, Kannada, Malayalam, etc.
        elif 0x0980 <= cp <= 0x0DFF:
            wide += 1
        # Thai
        elif 0x0E00 <= cp <= 0x0E7F:
            wide += 1
        # Korean Hangul
        elif 0xAC00 <= cp <= 0xD7AF:
            wide += 1
    if total == 0:
        return 0.6
    wide_ratio = wide / total
    # Blend: mostly wide chars → 0.85, mostly Latin → 0.6
    return 0.6 + wide_ratio * 0.25


def pt_emu(pt):
    return int(pt * PT_TO_EMU)


def set_section(section, width_pt, height_pt, left=28, right=28, top=20, bottom=14):
    section.page_width = pt_emu(width_pt)
    section.page_height = pt_emu(height_pt)
    section.left_margin = pt_emu(left)
    section.right_margin = pt_emu(right)
    section.top_margin = pt_emu(top)
    section.bottom_margin = pt_emu(bottom)
    section.header_distance = 0
    section.footer_distance = 0


def add_floating_image(paragraph, image_path, left_emu, top_emu, width_emu, height_emu,
                       behind=True):
    """Add a floating image anchored at absolute position."""
    run = paragraph.add_run()
    run.add_picture(image_path, width=width_emu, height=height_emu)

    drawing = run._element.findall(qn('w:drawing'))[0]
    inline = drawing.find(qn('wp:inline'))
    graphic = inline.find(qn('a:graphic'))
    extent = inline.find(qn('wp:extent'))
    docPr = inline.find(qn('wp:docPr'))
    cx, cy = extent.get('cx'), extent.get('cy')

    z_index = '0' if behind else '251660288'
    anchor = etree.SubElement(drawing, qn('wp:anchor'), {
        'distT': '0', 'distB': '0', 'distL': '0', 'distR': '0',
        'simplePos': '0', 'relativeHeight': z_index,
        'behindDoc': '1' if behind else '0',
        'locked': '0', 'layoutInCell': '1', 'allowOverlap': '1',
    })
    etree.SubElement(anchor, qn('wp:simplePos'), {'x': '0', 'y': '0'})
    posH = etree.SubElement(anchor, qn('wp:positionH'), {'relativeFrom': 'page'})
    etree.SubElement(posH, qn('wp:posOffset')).text = str(left_emu)
    posV = etree.SubElement(anchor, qn('wp:positionV'), {'relativeFrom': 'page'})
    etree.SubElement(posV, qn('wp:posOffset')).text = str(top_emu)
    etree.SubElement(anchor, qn('wp:extent'), {'cx': cx, 'cy': cy})
    etree.SubElement(anchor, qn('wp:effectExtent'), {'l': '0', 't': '0', 'r': '0', 'b': '0'})
    etree.SubElement(anchor, qn('wp:wrapNone'))
    anchor.append(docPr)
    etree.SubElement(anchor, qn('wp:cNvGraphicFramePr'))
    anchor.append(graphic)
    drawing.remove(inline)


def _build_run_element(span, parent):
    """Build a <w:r> element for a single span, appended to parent."""
    r_el = etree.SubElement(parent, qn('w:r'))
    rPr = etree.SubElement(r_el, qn('w:rPr'))

    font = _safe(span.get('font', 'Arial'))
    r_fonts = etree.SubElement(rPr, qn('w:rFonts'))
    r_fonts.set(qn('w:ascii'), font)
    r_fonts.set(qn('w:hAnsi'), font)
    r_fonts.set(qn('w:cs'), font)
    r_fonts.set(qn('w:eastAsia'), font)

    sz_val = str(int(span['size'] * 2))
    etree.SubElement(rPr, qn('w:sz')).set(qn('w:val'), sz_val)
    etree.SubElement(rPr, qn('w:szCs')).set(qn('w:val'), sz_val)

    if span.get('bold'):
        etree.SubElement(rPr, qn('w:b'))
    if span.get('italic'):
        etree.SubElement(rPr, qn('w:i'))

    color_el = etree.SubElement(rPr, qn('w:color'))
    color_el.set(qn('w:val'), _safe(span.get('color', '000000')))

    t_el = etree.SubElement(r_el, qn('w:t'))
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = _safe(span.get('text', ''))
    return r_el


def _is_paragraph_block(lines):
    """Check if a multi-line block should flow as a single paragraph.

    Returns True when all lines share a similar dominant font size,
    meaning they were merged from a continuous paragraph.
    """
    if len(lines) <= 1:
        return False
    sizes = []
    for line in lines:
        total_len = 0
        size_weight = {}
        for span in line.get('spans', []):
            slen = len(span.get('text', ''))
            total_len += slen
            sz = span.get('size', 11)
            size_weight[sz] = size_weight.get(sz, 0) + slen
        if size_weight:
            sizes.append(max(size_weight, key=size_weight.get))
    if not sizes:
        return False
    base = sizes[0]
    return all(abs(s - base) <= 0.5 for s in sizes)


def make_vml_textbox_run(lines, left_pt, top_pt, width_pt, height_pt,
                         flow_as_paragraph=False):
    """Build a VML textbox <w:r> element for a single text block.

    When flow_as_paragraph=True, all lines are merged into a single <w:p>
    with spaces between them, so Word wraps the text naturally within the
    fixed textbox width — matching browser CSS word-wrap behavior.
    """
    W = NSMAP['w']
    V = NSMAP['v']
    W10 = NSMAP['w10']

    r_root = etree.Element(_vqn('w:r'), nsmap=NSMAP)
    pict = etree.SubElement(r_root, _vqn('w:pict'))

    style = (
        f'position:absolute;'
        f'left:{left_pt:.2f}pt;top:{top_pt:.2f}pt;'
        f'width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;'
        f'z-index:251660288;'
        f'mso-position-horizontal-relative:page;'
        f'mso-position-vertical-relative:page;'
        f'mso-fit-shape-to-text:t'
    )
    shape = etree.SubElement(pict, _vqn('v:shape'), {
        'style': style, 'filled': 'f', 'stroked': 'f',
    })
    textbox = etree.SubElement(shape, _vqn('v:textbox'), {
        'inset': '0,0,0,0',
        'style': 'mso-fit-shape-to-text:t',
    })
    txbx = etree.SubElement(textbox, _vqn('w:txbxContent'))

    if flow_as_paragraph and len(lines) > 1:
        # Merge all lines into one <w:p> — Word wraps within textbox width
        p_el = etree.SubElement(txbx, _vqn('w:p'))
        pPr = etree.SubElement(p_el, _vqn('w:pPr'))
        spacing = etree.SubElement(pPr, _vqn('w:spacing'))
        spacing.set(_vqn('w:after'), '0')
        spacing.set(_vqn('w:before'), '0')
        spacing.set(_vqn('w:line'), '240')
        spacing.set(_vqn('w:lineRule'), 'auto')
        for li, line in enumerate(lines):
            # Add a space run between lines so words don't merge
            if li > 0:
                last_span = lines[li - 1].get('spans', [{}])[-1] if lines[li - 1].get('spans') else {}
                space_span = {
                    'text': ' ',
                    'font': last_span.get('font', 'Arial'),
                    'size': last_span.get('size', 11),
                    'color': last_span.get('color', '000000'),
                    'bold': last_span.get('bold', False),
                    'italic': last_span.get('italic', False),
                }
                _build_run_element(space_span, p_el)
            for span in line.get('spans', []):
                _build_run_element(span, p_el)
    else:
        # One <w:p> per line (original behavior for labels/headings)
        for line in lines:
            p_el = etree.SubElement(txbx, _vqn('w:p'))
            pPr = etree.SubElement(p_el, _vqn('w:pPr'))
            spacing = etree.SubElement(pPr, _vqn('w:spacing'))
            spacing.set(_vqn('w:after'), '0')
            spacing.set(_vqn('w:before'), '0')
            spacing.set(_vqn('w:line'), '240')
            spacing.set(_vqn('w:lineRule'), 'auto')
            for span in line.get('spans', []):
                _build_run_element(span, p_el)

    etree.SubElement(shape, _vqn('w10:wrap'), {'type': 'none'})
    return r_root


def add_positioned_table(anchor_p, table_data, page_height_pt):
    """Build a native Word table with absolute positioning, inserted after anchor_p.

    Uses tblpPr (table positioning properties) to place the table at the
    exact coordinates from the PDF. Handles col_span via gridSpan and uses
    actual cell bbox widths for accurate column sizing.
    """
    bbox = table_data['bbox']
    rows = table_data['rows']
    if not rows:
        return

    left_pt = bbox[0]
    top_pt = bbox[1]
    table_width_pt = bbox[2] - bbox[0]
    num_cols = table_data.get('num_cols', 0)

    if num_cols == 0:
        return

    # Determine grid column widths from cell bboxes.
    # Collect all column edge x-coordinates, then cluster nearby edges.
    raw_edges = set()
    raw_edges.add(round(bbox[0], 1))
    raw_edges.add(round(bbox[2], 1))
    for row in rows:
        for cell in row['cells']:
            cb = cell.get('bbox', [])
            if cb and len(cb) >= 4 and cb[2] > cb[0]:
                raw_edges.add(round(cb[0], 1))
                raw_edges.add(round(cb[2], 1))
    raw_edges = sorted(raw_edges)

    # Cluster edges within tolerance to avoid micro-columns
    EDGE_TOL = 3.0
    col_edges = []
    for e in raw_edges:
        if col_edges and abs(e - col_edges[-1]) < EDGE_TOL:
            col_edges[-1] = (col_edges[-1] + e) / 2
        else:
            col_edges.append(e)

    # Build grid columns from consecutive edges
    grid_widths_twips = []
    for j in range(len(col_edges) - 1):
        w = col_edges[j + 1] - col_edges[j]
        grid_widths_twips.append(int(w * 20))
    n_grid_cols = len(grid_widths_twips)

    if n_grid_cols == 0:
        return

    # Build table XML directly
    tbl_el = etree.SubElement(anchor_p._element.getparent(), qn('w:tbl'))
    anchor_p._element.addnext(tbl_el)

    # Table properties
    tbl_pr = etree.SubElement(tbl_el, qn('w:tblPr'))

    # Absolute positioning
    tblp_pr = etree.SubElement(tbl_pr, qn('w:tblpPr'))
    tblp_pr.set(qn('w:vertAnchor'), 'page')
    tblp_pr.set(qn('w:horzAnchor'), 'page')
    tblp_pr.set(qn('w:tblpX'), str(int(left_pt * 20)))
    tblp_pr.set(qn('w:tblpY'), str(int(top_pt * 20)))
    tblp_pr.set(qn('w:leftFromText'), '0')
    tblp_pr.set(qn('w:rightFromText'), '0')
    tblp_pr.set(qn('w:topFromText'), '0')
    tblp_pr.set(qn('w:bottomFromText'), '0')

    # Table width
    tbl_w = etree.SubElement(tbl_pr, qn('w:tblW'))
    tbl_w.set(qn('w:w'), str(int(table_width_pt * 20)))
    tbl_w.set(qn('w:type'), 'dxa')

    # Table layout fixed
    tbl_layout = etree.SubElement(tbl_pr, qn('w:tblLayout'))
    tbl_layout.set(qn('w:type'), 'fixed')

    # Borders
    tbl_borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = etree.SubElement(tbl_borders, qn(f'w:{border_name}'))
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), '999999')

    # Table grid
    tbl_grid = etree.SubElement(tbl_el, qn('w:tblGrid'))
    for gw in grid_widths_twips:
        grid_col = etree.SubElement(tbl_grid, qn('w:gridCol'))
        grid_col.set(qn('w:w'), str(gw))

    # Helper: find which grid column index a cell's left edge maps to
    def _grid_col_for_x(x_pt):
        x_r = round(x_pt, 1)
        best = 0
        best_dist = abs(col_edges[0] - x_r)
        for k, edge in enumerate(col_edges):
            d = abs(edge - x_r)
            if d < best_dist:
                best_dist = d
                best = k
        return best

    # Build rows and cells
    for ri, row in enumerate(rows):
        tr_el = etree.SubElement(tbl_el, qn('w:tr'))

        # Set fixed row height from cell bboxes to prevent table growth
        row_h_pt = 0
        for cell in row['cells']:
            cb = cell.get('bbox', [])
            if cb and len(cb) >= 4 and cb[3] > cb[1]:
                row_h_pt = max(row_h_pt, cb[3] - cb[1])
        if row_h_pt > 0:
            tr_pr = etree.SubElement(tr_el, qn('w:trPr'))
            tr_height = etree.SubElement(tr_pr, qn('w:trHeight'))
            tr_height.set(qn('w:val'), str(int(row_h_pt * 20)))
            tr_height.set(qn('w:hRule'), 'exact')

        grid_col_idx = 0  # track which grid column we're at

        for cell in row['cells']:
            cell_bbox = cell.get('bbox', [])
            col_span = cell.get('col_span', 1)

            # Calculate how many grid columns this cell spans from its bbox
            if cell_bbox and len(cell_bbox) >= 4 and cell_bbox[2] > cell_bbox[0]:
                start_gc = _grid_col_for_x(cell_bbox[0])
                end_gc = _grid_col_for_x(cell_bbox[2])
                actual_span = max(1, end_gc - start_gc)

                # Fill empty grid cells before this cell if there's a gap
                while grid_col_idx < start_gc and grid_col_idx < n_grid_cols:
                    empty_tc = etree.SubElement(tr_el, qn('w:tc'))
                    empty_pr = etree.SubElement(empty_tc, qn('w:tcPr'))
                    empty_w = etree.SubElement(empty_pr, qn('w:tcW'))
                    empty_w.set(qn('w:w'), str(grid_widths_twips[grid_col_idx]))
                    empty_w.set(qn('w:type'), 'dxa')
                    etree.SubElement(empty_tc, qn('w:p'))
                    grid_col_idx += 1

                cw = sum(grid_widths_twips[start_gc:end_gc]) if end_gc <= n_grid_cols else int((cell_bbox[2] - cell_bbox[0]) * 20)
            else:
                actual_span = col_span
                cw = sum(grid_widths_twips[grid_col_idx:grid_col_idx + actual_span]) if grid_col_idx + actual_span <= n_grid_cols else int(table_width_pt * 20 // n_grid_cols)

            tc_el = etree.SubElement(tr_el, qn('w:tc'))
            tc_pr = etree.SubElement(tc_el, qn('w:tcPr'))

            # Cell width
            tc_w = etree.SubElement(tc_pr, qn('w:tcW'))
            tc_w.set(qn('w:w'), str(cw))
            tc_w.set(qn('w:type'), 'dxa')

            # Grid span for merged cells
            if actual_span > 1:
                grid_span = etree.SubElement(tc_pr, qn('w:gridSpan'))
                grid_span.set(qn('w:val'), str(actual_span))

            # Cell margins (minimal to match PDF layout)
            tc_mar = etree.SubElement(tc_pr, qn('w:tcMar'))
            for side in ['top', 'left', 'bottom', 'right']:
                mar = etree.SubElement(tc_mar, qn(f'w:{side}'))
                mar.set(qn('w:w'), '0')
                mar.set(qn('w:type'), 'dxa')

            # Header shading
            if cell.get('is_header'):
                shading = etree.SubElement(tc_pr, qn('w:shd'))
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'E8E8E8')

            # Cell paragraph with text
            p_el = etree.SubElement(tc_el, qn('w:p'))
            p_pr = etree.SubElement(p_el, qn('w:pPr'))
            spacing = etree.SubElement(p_pr, qn('w:spacing'))
            spacing.set(qn('w:before'), '10')
            spacing.set(qn('w:after'), '10')
            spacing.set(qn('w:line'), '200')
            spacing.set(qn('w:lineRule'), 'atLeast')

            text = cell.get('text', '')
            if text:
                r_el = etree.SubElement(p_el, qn('w:r'))
                r_pr = etree.SubElement(r_el, qn('w:rPr'))
                sz = etree.SubElement(r_pr, qn('w:sz'))
                sz.set(qn('w:val'), '12')  # 6pt
                sz_cs = etree.SubElement(r_pr, qn('w:szCs'))
                sz_cs.set(qn('w:val'), '12')
                r_fonts = etree.SubElement(r_pr, qn('w:rFonts'))
                r_fonts.set(qn('w:ascii'), 'Arial')
                r_fonts.set(qn('w:hAnsi'), 'Arial')
                r_fonts.set(qn('w:cs'), 'Arial')
                r_fonts.set(qn('w:eastAsia'), 'Arial')
                if cell.get('is_header'):
                    etree.SubElement(r_pr, qn('w:b'))
                t_el = etree.SubElement(r_el, qn('w:t'))
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t_el.text = _safe(text)

            grid_col_idx += actual_span

        # Fill remaining grid columns in the row
        while grid_col_idx < n_grid_cols:
            empty_tc = etree.SubElement(tr_el, qn('w:tc'))
            empty_pr = etree.SubElement(empty_tc, qn('w:tcPr'))
            empty_w = etree.SubElement(empty_pr, qn('w:tcW'))
            empty_w.set(qn('w:w'), str(grid_widths_twips[grid_col_idx]))
            empty_w.set(qn('w:type'), 'dxa')
            etree.SubElement(empty_tc, qn('w:p'))
            grid_col_idx += 1


def build_docx(project_data, project_dir, output_path):
    """Build a .docx from structured page/block/image data.

    Args:
        project_data: dict with 'pages' list (each has blocks, images, dimensions)
        project_dir: base directory for resolving asset paths
        output_path: where to save the .docx
    """
    doc = Document()
    pages = project_data['pages']

    for i, page in enumerate(pages):
        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        set_section(section, page['width'], page['height'],
                    left=0, right=0, top=0, bottom=0)

        # Anchor paragraph for floating images
        anchor_p = doc.add_paragraph()
        anchor_p.paragraph_format.space_before = Pt(0)
        anchor_p.paragraph_format.space_after = Pt(0)

        # Place background image (text-stripped render)
        bg_path = os.path.join(project_dir, page['background_image'])
        if os.path.exists(bg_path):
            add_floating_image(
                anchor_p, bg_path,
                pt_emu(0), pt_emu(0),
                pt_emu(page['width']), pt_emu(page['height']),
                behind=True
            )

        # Place extracted images
        for img in page.get('images', []):
            img_path = os.path.join(project_dir, img['path'])
            if not os.path.exists(img_path):
                continue
            bbox = img['bbox']
            add_floating_image(
                anchor_p, img_path,
                pt_emu(bbox[0]), pt_emu(bbox[1]),
                pt_emu(bbox[2] - bbox[0]), pt_emu(bbox[3] - bbox[1]),
                behind=False
            )

        # Render table cell text as VML textboxes (same as regular text).
        # The background image already contains the table borders/headers,
        # so we just need text at the exact cell positions.
        page_tables = page.get('tables', [])
        table_bboxes = [t['bbox'] for t in page_tables]
        for table in page_tables:
            for row in table.get('rows', []):
                for cell in row.get('cells', []):
                    text = cell.get('text', '').strip()
                    if not text:
                        continue
                    cb = cell.get('bbox', [0, 0, 0, 0])
                    if cb[2] <= cb[0]:
                        continue
                    cw = cb[2] - cb[0]
                    ch = cb[3] - cb[1]
                    font_size = min(7, max(4, ch * 0.7))
                    is_header = cell.get('is_header', False)
                    cell_lines = [{'spans': [{
                        'text': text,
                        'font': 'Arial',
                        'size': round(font_size, 1),
                        'color': '000000',
                        'bold': is_header,
                        'italic': False,
                    }]}]
                    vml_run = make_vml_textbox_run(
                        cell_lines,
                        left_pt=cb[0], top_pt=cb[1],
                        width_pt=cw + 2, height_pt=ch,
                    )
                    anchor_p._element.append(vml_run)

        # Place text blocks as VML textboxes packed into the anchor paragraph
        # (packing into one paragraph prevents overflow onto extra pages)
        for block in page.get('blocks', []):
            bbox = block['bbox']
            if table_bboxes and _is_inside_table(bbox, table_bboxes):
                continue
            raw_w = bbox[2] - bbox[0]
            raw_h = bbox[3] - bbox[1]
            block_lines = block.get('lines', [])
            is_para = _is_paragraph_block(block_lines)

            if is_para:
                # Paragraph block: use bbox width as constraint, let Word wrap
                # (mirrors browser behavior: fixed-width container + word-wrap)
                width_pt = raw_w + 2
            else:
                # Single line / label: estimate width from text content
                max_line_width = 0
                for line in block_lines:
                    line_w = 0
                    for span in line.get('spans', []):
                        text = span.get('text', '')
                        char_w = span.get('size', 11) * _char_width_multiplier(text)
                        line_w += len(text) * char_w
                    max_line_width = max(max_line_width, line_w)
                width_pt = max(raw_w + 5, max_line_width + 10)

            vml_run = make_vml_textbox_run(
                block_lines,
                left_pt=bbox[0], top_pt=bbox[1],
                width_pt=width_pt,
                height_pt=raw_h * 1.2,
                flow_as_paragraph=is_para,
            )
            anchor_p._element.append(vml_run)

    doc.save(output_path)
    print(f"  [+] Built: {output_path}")
    return output_path
