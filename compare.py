"""
Compare extracted OCR text (.txt) against ground_truth.docx
Metrics: CAR (Character Accuracy Rate), WAR (Word Accuracy Rate),
         NumAcc (Numerical Accuracy), Layout score, Table score.

Usage:
  python compare.py                                  # auto-detect latest project
  python compare.py --project-id c4085b16            # specific project
  python compare.py --gt ground_truth.docx --txt extracted_text.txt
"""

import os
import re
import sys
import json
import argparse
from difflib import SequenceMatcher
from docx import Document
from lxml import etree


# ── text extraction helpers ──────────────────────────────────────────────────

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_V = "urn:schemas-microsoft-com:vml"


def extract_text_from_docx(doc_path):
    """Extract all <w:t> text from a .docx in reading order."""
    doc = Document(doc_path)
    body = doc.element.body
    parts = [t.text for t in body.iter(f"{{{NS_W}}}t") if t.text]
    return " ".join(parts)


def extract_text_from_txt(txt_path):
    """Read extracted_text.txt, stripping page markers and table markers."""
    with open(txt_path, encoding="utf-8") as f:
        raw = f.read()
    # Remove page markers and table markers for flat text comparison
    lines = []
    for line in raw.splitlines():
        if line.startswith("--- Page") or line.strip() in ("[TABLE]", "[/TABLE]"):
            continue
        if line.strip():
            lines.append(line.strip())
    return " ".join(lines)


def extract_tables_from_txt(txt_path):
    """Parse [TABLE]...[/TABLE] blocks from extracted_text.txt."""
    with open(txt_path, encoding="utf-8") as f:
        raw = f.read()
    tables = []
    current_table = None
    for line in raw.splitlines():
        stripped = line.strip()
        if stripped == "[TABLE]":
            current_table = []
        elif stripped == "[/TABLE]" and current_table is not None:
            tables.append(current_table)
            current_table = None
        elif current_table is not None and stripped:
            cells = [c.strip() for c in stripped.split(" | ")]
            current_table.append(cells)
    return tables


def extract_tables_from_docx(doc_path):
    """Extract tables from a .docx as list of row-lists."""
    doc = Document(doc_path)
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables


def extract_blocks_from_vml(doc_path):
    """Extract positioned text blocks from VML textboxes in a .docx."""
    doc = Document(doc_path)
    body = doc.element.body
    blocks = []
    for shape in body.iter(f"{{{NS_V}}}shape"):
        style = shape.get("style", "")
        pos = _parse_vml_style(style)
        texts = [t.text for t in shape.iter(f"{{{NS_W}}}t") if t.text]
        text = " ".join(texts).strip()
        if text:
            blocks.append({"text": text, **pos})
    return blocks


def extract_blocks_from_json(project_json_path):
    """Extract block positions from project.json for layout comparison."""
    with open(project_json_path, encoding="utf-8") as f:
        data = json.load(f)
    blocks = []
    for page in data.get("pages", []):
        for block in page.get("blocks", []):
            bbox = block.get("bbox", [0, 0, 0, 0])
            text_parts = []
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text_parts.append(span.get("text", ""))
            text = " ".join(text_parts).strip()
            if text and len(bbox) >= 4:
                blocks.append({
                    "text": text,
                    "left": bbox[0], "top": bbox[1],
                    "width": bbox[2] - bbox[0],
                    "height": bbox[3] - bbox[1],
                })
    return blocks


def _parse_vml_style(style):
    """Parse VML style string → {left, top, width, height} in pt."""
    result = {"left": 0, "top": 0, "width": 0, "height": 0}
    for part in style.split(";"):
        part = part.strip()
        for key in result:
            if part.startswith(f"{key}:"):
                val = part.split(":")[1].strip().replace("pt", "")
                try:
                    result[key] = float(val)
                except ValueError:
                    pass
    return result


# ── normalize / numbers ──────────────────────────────────────────────────────

def normalize(text):
    """Lowercase, collapse whitespace, strip punctuation."""
    text = text.lower()
    text = re.sub(r"[^\w\s]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_numbers(text):
    """Extract numeric tokens (integers / decimals) from text."""
    return re.findall(r"\d+(?:[.,]\d+)*", text)


# ── metrics ──────────────────────────────────────────────────────────────────

def calc_car(gt_text, out_text):
    """Character Accuracy Rate — longest-common-subsequence ratio."""
    gt_n = normalize(gt_text)
    out_n = normalize(out_text)
    if not gt_n:
        return 1.0 if not out_n else 0.0
    sm = SequenceMatcher(None, gt_n, out_n)
    matching = sum(block.size for block in sm.get_matching_blocks())
    return matching / max(len(gt_n), 1)


def calc_war(gt_text, out_text):
    """Word Accuracy Rate — matching words ratio."""
    gt_words = normalize(gt_text).split()
    out_words = normalize(out_text).split()
    if not gt_words:
        return 1.0 if not out_words else 0.0
    sm = SequenceMatcher(None, gt_words, out_words)
    matching = sum(block.size for block in sm.get_matching_blocks())
    return matching / max(len(gt_words), 1)


def calc_num_accuracy(gt_text, out_text):
    """Fraction of GT numbers found in output."""
    gt_nums = extract_numbers(gt_text)
    out_nums = extract_numbers(out_text)
    if not gt_nums:
        return 1.0
    out_set = set(out_nums)
    matched = sum(1 for n in gt_nums if n in out_set)
    return matched / len(gt_nums)


def calc_layout_score(json_blocks, docx_blocks, tolerance=15.0):
    """Layout accuracy: how well docx VML positions match project.json positions.

    Compares the block positions from project.json (source of truth for layout)
    against the positions in the generated .docx to verify the builder preserved
    layout correctly.
    """
    if not json_blocks and not docx_blocks:
        return 1.0
    if not json_blocks or not docx_blocks:
        return 0.0

    matched = 0
    used = set()
    for jb in json_blocks:
        best_dist = float("inf")
        best_idx = -1
        for i, db in enumerate(docx_blocks):
            if i in used:
                continue
            dist = ((jb["left"] - db["left"]) ** 2 + (jb["top"] - db["top"]) ** 2) ** 0.5
            if dist < best_dist:
                best_dist = dist
                best_idx = i
        if best_dist <= tolerance and best_idx >= 0:
            matched += 1
            used.add(best_idx)
    return matched / max(len(json_blocks), 1)


def calc_table_score(gt_tables, out_tables):
    """Cell-level text match between GT and output tables."""
    if not gt_tables and not out_tables:
        return 1.0
    if not gt_tables or not out_tables:
        # Both have tables but counts differ — score structure only
        if gt_tables and not out_tables:
            return 0.0
        if out_tables and not gt_tables:
            return 0.0

    total_cells = 0
    matched_cells = 0

    for gt_tbl, out_tbl in zip(gt_tables, out_tables):
        max_rows = max(len(gt_tbl), len(out_tbl))
        for ri in range(max_rows):
            gt_row = gt_tbl[ri] if ri < len(gt_tbl) else []
            out_row = out_tbl[ri] if ri < len(out_tbl) else []
            max_cols = max(len(gt_row), len(out_row))
            for ci in range(max_cols):
                total_cells += 1
                gt_cell = normalize(gt_row[ci]) if ci < len(gt_row) else ""
                out_cell = normalize(out_row[ci]) if ci < len(out_row) else ""
                if gt_cell == out_cell:
                    matched_cells += 1
                elif gt_cell and out_cell:
                    sm = SequenceMatcher(None, gt_cell, out_cell)
                    matched_cells += sm.ratio()

    return matched_cells / max(total_cells, 1)


# ── report ───────────────────────────────────────────────────────────────────

def compare(gt_text, out_text, gt_tables, out_tables, json_blocks, docx_blocks):
    """Run all comparisons and return metrics dict."""
    return {
        "CAR": calc_car(gt_text, out_text),
        "WAR": calc_war(gt_text, out_text),
        "NumAcc": calc_num_accuracy(gt_text, out_text),
        "Layout": calc_layout_score(json_blocks, docx_blocks),
        "Table": calc_table_score(gt_tables, out_tables),
    }


def detect_language_mismatch(gt_text, out_text):
    """Detect if GT and output use different scripts."""
    gt_non_ascii = sum(1 for c in gt_text if ord(c) > 127) / max(len(gt_text), 1)
    out_non_ascii = sum(1 for c in out_text if ord(c) > 127) / max(len(out_text), 1)
    return abs(gt_non_ascii - out_non_ascii) > 0.2


def print_report(metrics, gt_text, out_text):
    """Print a formatted comparison report."""
    print("=" * 60)
    print("  DOCX Comparison Report")
    print("=" * 60)
    print()

    if detect_language_mismatch(gt_text, out_text):
        print("  NOTE: Output appears to be in a different language/")
        print("  script. CAR/WAR compare extraction fidelity, not")
        print("  translation quality.")
        print()

    gt_nums = extract_numbers(gt_text)
    out_nums = extract_numbers(out_text)

    print(f"  GT chars : {len(normalize(gt_text)):>6}")
    print(f"  Out chars: {len(normalize(out_text)):>6}")
    print(f"  GT words : {len(normalize(gt_text).split()):>6}")
    print(f"  Out words: {len(normalize(out_text).split()):>6}")
    print(f"  GT nums  : {len(gt_nums):>6}  {gt_nums[:10]}{'...' if len(gt_nums) > 10 else ''}")
    print(f"  Out nums : {len(out_nums):>6}  {out_nums[:10]}{'...' if len(out_nums) > 10 else ''}")
    print()
    print("-" * 60)
    print(f"  {'Metric':<20} {'Score':>10} {'Pct':>8}")
    print("-" * 60)
    for name, val in metrics.items():
        bar = "#" * int(val * 30)
        print(f"  {name:<20} {val:>10.4f} {val*100:>7.1f}%  |{bar}")
    print("-" * 60)
    avg = sum(metrics.values()) / len(metrics)
    print(f"  {'AVERAGE':<20} {avg:>10.4f} {avg*100:>7.1f}%")
    print("=" * 60)


# ── main ─────────────────────────────────────────────────────────────────────

def find_project_dir():
    """Auto-detect the most recent project directory."""
    projects_dir = "projects"
    if not os.path.isdir(projects_dir):
        return None
    dirs = sorted(
        os.listdir(projects_dir),
        key=lambda d: os.path.getmtime(os.path.join(projects_dir, d)),
        reverse=True,
    )
    for d in dirs:
        pdir = os.path.join(projects_dir, d)
        if os.path.isdir(pdir) and os.path.exists(os.path.join(pdir, "project.json")):
            return pdir
    return None


def main():
    parser = argparse.ArgumentParser(description="Compare extracted text against ground truth")
    parser.add_argument("--gt", default="ground_truth.docx",
                        help="Ground truth .docx path")
    parser.add_argument("--txt", default=None,
                        help="Extracted text .txt path (from download)")
    parser.add_argument("--docx", default=None,
                        help="Output .docx path (for layout comparison)")
    parser.add_argument("--project-id", default=None,
                        help="Project ID (auto-resolves paths)")
    parser.add_argument("--project-json", default=None,
                        help="project.json path (for layout/table data)")
    args = parser.parse_args()

    # Resolve project directory
    if args.project_id:
        project_dir = os.path.join("projects", args.project_id)
    elif not args.txt:
        project_dir = find_project_dir()
    else:
        project_dir = None

    # Resolve file paths
    txt_path = args.txt
    docx_path = args.docx
    json_path = args.project_json

    if project_dir:
        if not txt_path:
            txt_path = os.path.join(project_dir, "extracted_text.txt")
        if not docx_path:
            for name in ("output.docx", "test_output.docx"):
                candidate = os.path.join(project_dir, name)
                if os.path.exists(candidate):
                    docx_path = candidate
                    break
        if not json_path:
            json_path = os.path.join(project_dir, "project.json")

    # Validate
    if not txt_path or not os.path.exists(txt_path):
        # Generate the txt from project.json if it doesn't exist yet
        if json_path and os.path.exists(json_path):
            print(f"  extracted_text.txt not found, generating from project.json...")
            txt_path = _generate_txt_from_json(json_path, project_dir)
        else:
            print("  ERROR: No extracted_text.txt found. Run a download first,")
            print("  or specify --txt path.")
            sys.exit(1)

    if not os.path.exists(args.gt):
        print(f"  ERROR: Ground truth not found: {args.gt}")
        sys.exit(1)

    print(f"  Ground truth : {args.gt}")
    print(f"  Extracted txt: {txt_path}")
    if docx_path:
        print(f"  Output docx  : {docx_path}")
    if json_path:
        print(f"  Project JSON : {json_path}")
    print()

    # ── Extract data ──
    gt_text = extract_text_from_docx(args.gt)
    out_text = extract_text_from_txt(txt_path)

    # Tables: from GT docx and from extracted txt
    gt_tables = extract_tables_from_docx(args.gt)
    out_tables = extract_tables_from_txt(txt_path)

    # Layout: compare project.json positions vs docx VML positions
    json_blocks = []
    docx_blocks = []
    if json_path and os.path.exists(json_path):
        json_blocks = extract_blocks_from_json(json_path)
    if docx_path and os.path.exists(docx_path):
        docx_blocks = extract_blocks_from_vml(docx_path)

    # ── Compare ──
    metrics = compare(gt_text, out_text, gt_tables, out_tables, json_blocks, docx_blocks)
    print_report(metrics, gt_text, out_text)


def _generate_txt_from_json(json_path, project_dir):
    """Generate extracted_text.txt from project.json (same logic as download router)."""
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    lines = []
    for page in data.get("pages", []):
        page_num = page.get("page", 0)
        lines.append(f"--- Page {page_num + 1} ---")
        for block in page.get("blocks", []):
            block_texts = []
            for line in block.get("lines", []):
                span_texts = [s.get("text", "") for s in line.get("spans", [])]
                block_texts.append("".join(span_texts))
            text = " ".join(block_texts).strip()
            if text:
                lines.append(text)
        for table in page.get("tables", []):
            lines.append("[TABLE]")
            for row in table.get("rows", []):
                cells = [c.get("text", "") for c in row.get("cells", [])]
                lines.append(" | ".join(cells))
            lines.append("[/TABLE]")
        lines.append("")

    txt_path = os.path.join(project_dir or ".", "extracted_text.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"  [+] Generated: {txt_path}")
    return txt_path


if __name__ == "__main__":
    main()
